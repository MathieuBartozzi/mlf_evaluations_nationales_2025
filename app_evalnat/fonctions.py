import fitz
import streamlit as st
from openai import OpenAI
import markdown
import io
from weasyprint import HTML
import os
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from io import BytesIO
import plotly.io as pio
from fpdf import FPDF
from fonctions_viz import *

import imgkit
import tempfile

def fig_to_png(fig):
    # Export de la figure Plotly en HTML (sans Kaleido)
    html_str = pio.to_html(fig, full_html=False)

    # Sauvegarde HTML temporaire
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as f_html:
        f_html.write(html_str.encode("utf-8"))
        html_path = f_html.name

    # Conversion HTML → PNG via wkhtmltoimage (indépendant de Chrome)
    png_bytes = imgkit.from_file(html_path, False)  # False = retourne bytes

    return png_bytes

# ---------------------------------------------------------
#   FONCTION  : Générer un rapport d'analyse pour un établissement
# ---------------------------------------------------------
def generer_rapport_etablissement(df, selected_etablissement, contexte_local=None, pdf_text=None):
    """
    Génère un rapport d'analyse pédagogique pour un établissement
    en intégrant Matière, Domaine et Compétence.
    """

    # 🔎 Informations générales
    ville, pays = df.iloc[0][["Ville", "Pays"]]
    niveaux = ", ".join(df["Niveau"].unique())

    # 📊 Résultats moyens hiérarchisés
    resultats = (
        df.groupby(["Matière", "Domaine", "Compétence", "Niveau"])["Valeur"]
        .mean()
        .reset_index()
    )

    # Reformater en texte structuré
    resultats_str = ""

    for matiere in resultats["Matière"].unique():
        resultats_str += f"\n## 📘 {matiere}\n"
        sous_df_mat = resultats[resultats["Matière"] == matiere]

        for domaine in sous_df_mat["Domaine"].unique():
            resultats_str += f"\n### 🔹 Domaine : {domaine}\n"
            sous_df_dom = sous_df_mat[sous_df_mat["Domaine"] == domaine]

            for _, row in sous_df_dom.iterrows():
                resultats_str += (
                    f"- **{row['Niveau']} - {row['Compétence']}** : "
                    f"{row['Valeur']:.1f}%\n"
                )

    # 📝 Titre + avertissement
    titre_rapport = (
        f"Rapport d'analyse pour l'établissement {selected_etablissement} "
        f"({ville}, {pays})\nDonnées des évaluations nationales 2025"
    )

# Construction du prompt OpenAI
    prompt = f"""

    Tu es un expert en éducation et en analyse de données scolaires.
    Ton rôle : aider un chef d'établissement à piloter la pédagogie via une analyse de données factuelle et bienveillante.

    TON OBJECTIF DE FORME (TRES IMPORTANT POUR L'EXPORT PDF) :
    1. Utilise strictement la syntaxe Markdown.
    2. Pour les titres, utilise seulement des niveaux ### (H3) et #### (H4). Ne jamais utiliser de # (H1) et ## (H2)
    3. IMPORTANT : Saute TOUJOURS une ligne vide avant de commencer une liste à puces.
    4. LISTES À PUCES :
       - Utilise des tirets "-" pour le premier niveau.
       - IMPORTANT : Pour les détails ou sous-points, crée une SOUS-LISTE en indentant (décalant) de 4 espaces.
       - Exemple :
         - Point principal
             - Détail du point principal (indenté)
    5. Mets en **gras** les concepts clés ou les chiffres importants.

    TON OBJECTIF DE FOND :
    - Ton ton doit être professionnel, neutre mais constructif.
    - Les constats chiffrés sont à l'indicatif (c'est factuel).
    - Les pistes d'action sont au conditionnel (tu es conseiller).
    - Ne termine pas par une formule de politesse, une signature ou une question. Finis directement après la dernière partie.

    ---

    # {titre_rapport}

    ### **Contexte**
    L'établissement **{selected_etablissement}**, situé à **{ville}, {pays}**, a récemment obtenu des résultats aux évaluations nationales pour les niveaux suivants : **{niveaux}**.

    **Scores moyens par niveau et par compétence :**
    {resultats_str}

    """

    if contexte_local:
        prompt += f"\n**Informations spécifiques fournies par la direction d'ecole :**\n{contexte_local}\n"

    # Ajouter le contenu extrait du PDF si disponible
    if pdf_text:
        prompt += f"\n**Informations complémentaires extraites du document joint :**\n{pdf_text[:1500]}..."  # Limite à 1500 caractères pour éviter un prompt trop long



    # 🧱 STRUCTURE D'ANALYSE DEMANDÉE
    prompt += """
---

Rédige le rapport en suivant strictement ce plan :

## Analyse des résultats

### 1. Tendances marquantes par niveau
*Sous-titres suggérés : Forces, Fragilités, Écarts*
- Pour chaque niveaux (CP, CE1, CE2, CM1, CM2) :
    - Présente les forces et réussites observées.
    - Identifie les domaines ou compétences qui semblent fragiles.
- Mentionne les éventuelles ruptures ou écarts significatifs.

### 2. Interprétation pédagogique
- Formule des hypothèses explicatives (ex: impact de la phonologie sur la lecture, automatismes...).
- Fais le lien entre les compétences (ex: le vocabulaire impacte-t-il la résolution de problèmes ?).
- Signale les compétences transversales qui pourraient jouer un rôle.

### 3. Pistes d'amélioration
- Propose des stratégies concrètes (rituels, différenciation, co-intervention...).
- Suggère des dispositifs spécifiques (APC, stages...).

### 4. Besoins de formation pour les enseignants
- Suggère des thématiques de formation pour les enseignants (didactique, gestion de classe, etc.).
- Propose des modalités (formation établissement, auto-formation, pairs...).
"""

    # 🧠 Appel API OpenAI
    client = OpenAI(api_key=st.secrets['openai']["OPENAI_API_KEY"])

    try:
        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[{"role": "user", "content": prompt}],
        )
        return response.choices[0].message.content, None

    except Exception as e:
        return None, f"Erreur lors de la génération du rapport : {e}"




# ---------------------------------------------------------
#   INTERFACE STREAMLIT : Bouton + affichage du rapport
# ---------------------------------------------------------


# Fonction pour extraire un texte limité à 3 pages
def extract_text_from_pdf(pdf_file, max_pages=3):
    """Extrait le texte des X premières pages d'un PDF, avec une limite sur le nombre de mots."""
    text = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break  # Stop après le nombre de pages défini
            text += page.get_text("text") + "\n"
    return text.strip()


def export_markdown_to_pdf(md_text, pdf_path="/tmp/rapport.pdf"):
    # Convert Markdown → HTML
    html = markdown.markdown(md_text)

    # Convert HTML → PDF
    HTML(string=html).write_pdf(pdf_path)

    return pdf_path


@st.cache_data(show_spinner=False)
def convert_to_pdf_data(md_text, ecole_nom):
    """
    Convertit le Markdown en PDF avec gestion des sous-listes (puces vides) et Logo.
    """

    # --- 0. Préparation du LOGO (Code inchangé) ---
    logo_html = ""
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(current_dir, "logo_mlfmonde.png")
        if not os.path.exists(logo_path):
             logo_path = os.path.join(current_dir, "..", "logo_mlfmonde.png")

        with open(logo_path, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode()
        logo_html = f'<img src="data:image/png;base64,{encoded_string}" class="logo">'
    except Exception:
        logo_html = ""

    # --- 1. Conversion Markdown -> HTML ---
    html_content = markdown.markdown(md_text, extensions=['tables', 'smarty'])

    # --- 2. CSS Professionnel Amélioré (Sous-listes) ---
    css_style = f"""
    <style>
        @page {{
            size: A4;
            margin: 2.5cm 2cm;
            @top-center {{
                content: "Rapport d'analyse - {ecole_nom}";
                font-family: 'Helvetica', sans-serif;
                font-size: 8pt;
                color: #aaa;
            }}
            @bottom-right {{
                content: "Page " counter(page);
                font-family: 'Helvetica', sans-serif;
                font-size: 9pt;
            }}
        }}

        body {{
            font-family: 'Helvetica', 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #2c3e50;
            text-align: justify;
        }}

        img.logo {{ display: block; margin: 0 auto 20px auto; width: 150px; }}

        h1 {{
            color: #0e4e7e; font-size: 22pt; border-bottom: 3px solid #0e4e7e;
            padding-bottom: 15px; margin-bottom: 30px; text-align: center; margin-top: 0;
        }}

        h2 {{
            color: #0e4e7e; font-size: 16pt; margin-top: 25px; margin-bottom: 15px;
            border-left: 5px solid #0e4e7e; padding-left: 15px; background-color: #f4f8fb;
            padding-top: 5px; padding-bottom: 5px;
        }}

        h3 {{
            color: #d35400; font-size: 13pt; margin-top: 20px; margin-bottom: 10px;
            font-weight: bold; text-transform: uppercase;
        }}

        /* --- GESTION DES LISTES IMBRIQUÉES --- */

        /* Niveau 1 : Puces pleines */
        ul {{
            margin-top: 5px;
            margin-bottom: 15px;
            padding-left: 20px;
            list-style-type: disc; /* Puce pleine par défaut */
        }}

        /* Niveau 2 : Listes dans des listes (Sous-points) */
        ul ul {{
            margin-top: 5px;
            margin-bottom: 5px;
            list-style-type: circle; /* CERCLE VIDE (Comme sur votre image) */
            color: #555; /* Légèrement plus gris pour la hiérarchie */
        }}

        /* Espacement des éléments */
        li {{
            margin-bottom: 6px;
            padding-left: 5px;
        }}

        /* Si numérotation (1. 2. 3.) */
        ol {{ padding-left: 20px; margin-bottom: 15px; }}
        ol li {{ margin-bottom: 8px; }}

        strong {{ color: #2c3e50; font-weight: 700; }}

        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; font-size: 10pt; }}
        th {{ background-color: #0e4e7e; color: white; padding: 8px; text-align: left; }}
        td {{ border: 1px solid #ddd; padding: 8px; }}
    </style>
    """

    # --- 3. Assemblage ---
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head><meta charset="utf-8">{css_style}</head>
    <body>
        {logo_html}
        <h1>Rapport d'analyse : {ecole_nom}</h1>
        {html_content}
        <br><br><hr style="border: 0; border-top: 1px solid #eee;">
        <p style="font-size: 9pt; color: #888; text-align: center; font-style: italic;">
            Document généré automatiquement par une IA.
        </p>
    </body>
    </html>
    """

    # --- 4. Génération PDF ---
    pdf_file = io.BytesIO()
    HTML(string=full_html).write_pdf(pdf_file)
    pdf_file.seek(0)
    return pdf_file.getvalue()


@st.cache_data(show_spinner=False)
def convert_to_word_data(md_text, ecole_nom):
    """
    Génère un fichier Word (.docx) avec le style du rapport (Logo, Couleurs, Hiérarchie).
    """
    doc = Document()

    # --- A. CONFIGURATION DES STYLES (Pour imiter le CSS du PDF) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Couleur Bleu MLF (#0e4e7e -> RGB 14, 78, 126)
    bleu_mlf = RGBColor(14, 78, 126)
    # Couleur Orange (#d35400 -> RGB 211, 84, 0)
    orange_mlf = RGBColor(211, 84, 0)

    # --- B. AJOUT DU LOGO ---
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(current_dir, "logo_mlfmonde.png")
        if not os.path.exists(logo_path):
             logo_path = os.path.join(current_dir, "..", "logo_mlfmonde.png")

        # On ajoute l'image centrée
        doc.add_picture(logo_path, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass # Pas de logo, pas grave

    # --- C. TITRE PRINCIPAL ---
    titre = doc.add_heading(f"Rapport d'analyse : {ecole_nom}", level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titre = titre.runs[0]
    run_titre.font.color.rgb = bleu_mlf
    run_titre.font.name = 'Arial'
    run_titre.font.bold = True

    # --- D. PARSING DU MARKDOWN ---
    # On lit le texte ligne par ligne pour appliquer les styles
    lines = md_text.split('\n')

    for line in lines:
        line = line.strip('\r') # Nettoyage

        # 1. Titres H2 (##) -> Bleu avec ligne
        if line.startswith("## "):
            text = line.replace("## ", "").strip()
            # On utilise le style Heading 1 de Word mais on le customise
            h = doc.add_heading(text.upper(), level=1)
            run = h.runs[0]
            run.font.color.rgb = bleu_mlf
            run.font.size = Pt(14)
            run.font.name = 'Arial'

        # 2. Titres H3 (###) -> Orange
        elif line.startswith("### "):
            text = line.replace("### ", "").strip()
            h = doc.add_heading(text, level=2)
            run = h.runs[0]
            run.font.color.rgb = orange_mlf
            run.font.size = Pt(12)
            run.font.name = 'Arial'

        # 3. Listes à puces (Niveau 1)
        elif line.strip().startswith("- "):
            text = line.strip().replace("- ", "")
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, text) # Gestion du gras **text**

        # 4. Listes à puces (Niveau 2 - Indenté)
        elif line.strip().startswith("- ") and (line.startswith("    ") or line.startswith("\t")):
            # C'est une sous-puce
            text = line.strip().replace("- ", "")
            p = doc.add_paragraph(style='List Bullet 2')
            _add_rich_text(p, text)

        # 5. Paragraphes normaux (si ligne non vide)
        elif line.strip():
            p = doc.add_paragraph()
            _add_rich_text(p, line)
            p.paragraph_format.space_after = Pt(8)

    # --- E. SAUVEGARDE EN MÉMOIRE ---
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def _add_rich_text(paragraph, text):
    """
    Petite fonction utilitaire pour gérer le **gras** dans Word
    """
    # On découpe le texte par les marqueurs **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # C'est du gras
            clean_text = part.replace("**", "")
            run = paragraph.add_run(clean_text)
            run.font.bold = True
        else:
            # Texte normal
            paragraph.add_run(part)



def generate_pdf(df_ecole, df_global, df_feat, df_pca,
                 ecole_selectionnee, ordre_niveaux, palette):
    """
    Génère un PDF complet et mis en forme, reproduisant la page Streamlit :
    - Carte d'identité
    - Métriques
    - Radar
    - Scatter
    - Heatmaps FR & MATH
    - Courbe progression
    - Profilage (PCA + clusters)
    - Recommandations
    """

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)

    # --------------------------------------------------
    # 0. Pré-calculs
    # --------------------------------------------------
    info_ecole = df_ecole[["Réseau", "Statut", "Homologué"]].drop_duplicates().iloc[0]

    moy_gen,  delta_gen  = get_moyenne_et_delta(df_global, df_ecole)
    moy_fr,   delta_fr   = get_moyenne_et_delta(df_global, df_ecole, "Français")
    moy_math, delta_math = get_moyenne_et_delta(df_global, df_ecole, "Mathématiques")

    cluster_id = int(df_feat.loc[ecole_selectionnee, "cluster"])
    profil = cluster_id + 1

    pc1 = df_pca.loc[df_pca["Nom_ecole"] == ecole_selectionnee, "PC1"].values[0]
    pc2 = df_pca.loc[df_pca["Nom_ecole"] == ecole_selectionnee, "PC2"].values[0]
    pc3 = df_pca.loc[df_pca["Nom_ecole"] == ecole_selectionnee, "PC3"].values[0]

    recommandations = get_recommandations_profil(profil)

    # --------------------------------------------------
    # PAGE 1 : TITRE + CARTE IDENTITÉ + MÉTRIQUES + RADAR + SCATTER
    # --------------------------------------------------
    pdf.add_page()

    # Titre
    pdf.set_font("Arial", "B", 18)
    pdf.cell(0, 10, f"Rapport - {ecole_selectionnee}", ln=True)

    # Sous-titre
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(
        0, 7,
        "Synthèse visuelle des performances et du profil global de l'établissement "
        "au regard des évaluations nationales."
    )
    pdf.ln(4)

    # Carte d'identité
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 8, "Carte d'identité de l'établissement", ln=True)

    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(
        0, 6,
        f"- Réseau : {info_ecole['Réseau']}\n"
        f"- Statut : {info_ecole['Statut']}\n"
        f"- Homologué : {info_ecole['Homologué']}"
    )
    pdf.ln(4)

    # Métriques
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 8, "Résultats globaux", ln=True)

    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(
        0, 6,
        f"- Moyenne générale : {moy_gen:.1f} %  (écart réseau : {delta_gen:+.1f} pts)\n"
        f"- Français : {moy_fr:.1f} %  (écart réseau : {delta_fr:+1.1f} pts)\n"
        f"- Mathématiques : {moy_math:.1f} %  (écart réseau : {delta_math:+1.1f} pts)"
    )
    pdf.ln(6)

    # ---------------------------- Radar ----------------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Positionnement par domaine", ln=True)

    fig_radar = plot_radar_domaine(df_ecole, df_global, ecole_selectionnee, palette, return_fig=True)
    # img_radar = pio.to_image(fig_radar, format="png", scale=3)
    img_radar = fig_to_png(fig_radar)
    pdf.image(BytesIO(img_radar), w=170)
    pdf.ln(4)

    # -------------------------- Scatter ----------------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Comparaison établissement / réseau", ln=True)

    fig_scatter = plot_scatter_comparatif(df_global, ecole_selectionnee, palette, return_fig=True)
    img_scatter = pio.to_image(fig_scatter, format="png", scale=3)
    pdf.image(BytesIO(img_scatter), w=170)

    # --------------------------------------------------
    # PAGE 2 : HEATMAP FR + MATH + COURBE PROGRESSION
    # --------------------------------------------------
    pdf.add_page()

    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Progression des apprentissages de CP à CM2", ln=True)

    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(
        0, 6,
        "Les graphiques suivants présentent la progression des compétences et "
        "la régularité des apprentissages dans le temps."
    )
    pdf.ln(4)

    # ---------------------- Heatmap Français ----------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Heatmap des compétences - Français", ln=True)

    fig_h_fr = plot_heatmap_competences(df_ecole, "Français", ordre_niveaux, return_fig=True)
    img_h_fr = pio.to_image(fig_h_fr, format="png", scale=3)
    pdf.image(BytesIO(img_h_fr), w=170)
    pdf.ln(4)

    # ---------------------- Heatmap Mathématiques ----------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Heatmap des compétences - Mathématiques", ln=True)

    fig_h_math = plot_heatmap_competences(df_ecole, "Mathématiques", ordre_niveaux, return_fig=True)
    img_h_math = pio.to_image(fig_h_math, format="png", scale=3)
    pdf.image(BytesIO(img_h_math), w=170)
    pdf.ln(4)

    # ---------------------- Courbe progression ----------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Courbe des moyennes par niveau", ln=True)

    fig_line = plot_line_chart(df_ecole, palette, ordre_niveaux, return_fig=True)
    img_line = pio.to_image(fig_line, format="png", scale=3)
    pdf.image(BytesIO(img_line), w=170)

    # --------------------------------------------------
    # PAGE 3 : PROFILAGE + PCA + CLUSTERS + AXES + RECO
    # --------------------------------------------------
    pdf.add_page()

    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Profil de l'établissement", ln=True)
    pdf.ln(3)

    # Profil
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(
        0, 6,
        f"L'établissement appartient au **Profil {profil}** dans la classification "
        f"issue de l'analyse en composantes principales."
    )
    pdf.ln(4)

    # ---------------------- Camembert clusters ----------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Répartition des profils dans le réseau", ln=True)

    fig_pie = plot_pie_clusters(df_feat, return_fig=True)
    img_pie = pio.to_image(fig_pie, format="png", scale=3)
    pdf.image(BytesIO(img_pie), w=120)
    pdf.ln(4)

    # ---------------------- PCA 3D ----------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Position de l'établissement dans l'espace des profils (PCA)", ln=True)

    fig_pca = plot_pca_3d(df_pca, ecole_selectionnee, palette, return_fig=True)
    img_pca = pio.to_image(fig_pca, format="png", scale=3)
    pdf.image(BytesIO(img_pca), w=170)
    pdf.ln(4)

    # ---------------------- Axes PCA ----------------------
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 7, "Interprétation des axes", ln=True)

    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(
        0, 6,
        f"Axe 1 - Fondamentaux : {pc1:.2f}\n"
        f"Axe 2 - Automatisation : {pc2:.2f}\n"
        f"Axe 3 - Complexité : {pc3:.2f}\n\n"
        "Les valeurs positives indiquent un positionnement supérieur à la moyenne du réseau, "
        "les valeurs négatives suggèrent un besoin de renforcement."
    )

    # ---------------------- Recommandations ----------------------
    pdf.add_page()
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, f"Recommandations pédagogiques - Profil {profil}", ln=True)

    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 6, recommandations)

    # --------------------------------------------------
    # EXPORT
    # --------------------------------------------------
    pdf_bytes = bytes(pdf.output(dest="S"))
    return pdf_bytes
